#!/usr/bin/env python3
"""
Comprehensive test suite for ArchMCP - 30 tests across both modes
Generates Word document with results, text, and images
"""
import sys
from pathlib import Path

# Add ArchMCP-Common to path
sys.path.insert(0, str(Path(__file__).parent.parent.parent / 'ArchMCP-Common'))

from bedrock_analyzer import BedrockAnalyzer
from enhanced_search import EnhancedSearch
from docx import Document
from docx.shared import Inches, Pt
from datetime import datetime
import yaml

# Load config
config_path = Path(__file__).parent.parent.parent / 'ArchMCP-Common' / 'config' / 'bedrock_config.yaml'
with open(config_path, 'r') as f:
    config = yaml.safe_load(f)

bedrock_config = config.get('bedrock', {})
analyzer = BedrockAnalyzer(
    region_name=bedrock_config.get('region', 'us-east-1'),
    model_id=bedrock_config.get('model_id', 'us.anthropic.claude-sonnet-4-20250514-v1:0')
)
search = EnhancedSearch()
top_k_icons = bedrock_config.get('top_k_icons', 3)

# Test cases
KEYWORD_TESTS = [
    "EC2",
    "S3",
    "Lambda",
    "RDS, DynamoDB",
    "VPC, subnet",
    "CloudFront, Route53",
    "ECS, EKS",
    "API Gateway",
    "SQS, SNS",
    "CloudWatch",
    "IAM, Cognito",
    "ElastiCache",
    "Kinesis",
    "Step Functions",
    "Glue, Athena"
]

DESCRIPTION_TESTS = [
    "A simple web application with EC2 instances and S3 storage",
    "Serverless API using Lambda and API Gateway with DynamoDB database",
    "Microservices architecture with ECS containers behind a load balancer",
    "Data pipeline with Kinesis streams, Lambda processing, and S3 storage",
    "Machine learning workflow using SageMaker and S3",
    "Multi-tier application with EC2, RDS, and ElastiCache",
    "Event-driven architecture with EventBridge, Lambda, and SQS",
    "Static website hosted on S3 with CloudFront CDN",
    "Container orchestration using EKS with Auto Scaling",
    "Real-time analytics with Kinesis Data Analytics and Redshift",
    "Batch processing with AWS Batch and S3",
    "IoT solution with IoT Core, Lambda, and DynamoDB",
    "Content delivery with CloudFront and S3 origin",
    "Hybrid cloud with Direct Connect and VPN",
    "Disaster recovery with backup to S3 Glacier"
]

def run_keyword_test(keyword):
    """Run a single keyword test"""
    keywords = [k.strip() for k in keyword.split(',')]
    results = []
    
    for kw in keywords:
        available_icons = list(search.keywords_mapping.keys())
        selected = analyzer.select_icons_for_keyword(kw, available_icons, top_n=top_k_icons)
        results.extend(selected)
    
    return list(set(results))  # Remove duplicates

def run_description_test(description):
    """Run a single description test"""
    services = analyzer.analyze_text_description(description, mode='local')
    icons = []
    
    for service in services:
        matches = search.find_service_matches(service, {})
        if matches:
            # Take top K matches per service
            for match in matches[:top_k_icons]:
                icons.append(match['name'])
    
    return list(set(icons))  # Remove duplicates

def add_icon_images(doc, icon_names):
    """Add icon images to document"""
    table = doc.add_table(rows=1, cols=min(len(icon_names), 4))
    cells = table.rows[0].cells
    
    for idx, icon_name in enumerate(icon_names[:4]):  # Max 4 images per row
        if icon_name in search.keywords_mapping:
            mapping = search.keywords_mapping[icon_name]
            
            # Get icon path
            if mapping.get('path'):
                icon_path = Path(mapping['path'])
            else:
                page = mapping.get('page', 0)
                icon_path = Path(f'icons/page{page}_icons/{icon_name}')
            
            if icon_path.exists():
                try:
                    paragraph = cells[idx].paragraphs[0]
                    run = paragraph.add_run()
                    run.add_picture(str(icon_path), width=Inches(1.0))
                    # Add icon name below
                    cells[idx].add_paragraph(icon_name.replace('.png', '').replace('_', ' '))
                except Exception as e:
                    cells[idx].text = f"Error: {icon_name}"

def main():
    doc = Document()
    
    # Title
    title = doc.add_heading('ArchMCP Comprehensive Test Results', 0)
    doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph(f'Total Tests: 30 (15 keywords + 15 descriptions)')
    doc.add_paragraph('')
    
    print("🧪 Running 30 comprehensive tests...")
    print("=" * 60)
    
    # KEYWORD TESTS
    doc.add_heading('Part 1: Keywords Mode Tests (15 tests)', 1)
    doc.add_paragraph('Testing Bedrock-powered intelligent icon selection for keywords')
    doc.add_paragraph('')
    
    for idx, keyword in enumerate(KEYWORD_TESTS, 1):
        print(f"\n[{idx}/15] Keywords Test: {keyword}")
        
        try:
            icons = run_keyword_test(keyword)
            
            # Add to document
            doc.add_heading(f'Test {idx}: "{keyword}"', 2)
            doc.add_paragraph(f'Mode: Keywords')
            doc.add_paragraph(f'Input: {keyword}')
            doc.add_paragraph(f'Icons Found: {len(icons)}')
            doc.add_paragraph(f'Results: {", ".join(icons)}')
            
            # Add images
            if icons:
                add_icon_images(doc, icons)
            
            doc.add_paragraph('')
            print(f"   ✅ Found {len(icons)} icons: {', '.join(icons[:3])}{'...' if len(icons) > 3 else ''}")
            
        except Exception as e:
            doc.add_paragraph(f'❌ Error: {str(e)}')
            print(f"   ❌ Error: {e}")
    
    # DESCRIPTION TESTS
    doc.add_page_break()
    doc.add_heading('Part 2: Description Mode Tests (15 tests)', 1)
    doc.add_paragraph('Testing Bedrock AI analysis of architecture descriptions')
    doc.add_paragraph('')
    
    for idx, description in enumerate(DESCRIPTION_TESTS, 1):
        print(f"\n[{idx}/15] Description Test: {description[:50]}...")
        
        try:
            icons = run_description_test(description)
            
            # Add to document
            doc.add_heading(f'Test {idx + 15}: Architecture Description', 2)
            doc.add_paragraph(f'Mode: Description')
            doc.add_paragraph(f'Input: {description}')
            doc.add_paragraph(f'Icons Found: {len(icons)}')
            doc.add_paragraph(f'Results: {", ".join(icons)}')
            
            # Add images
            if icons:
                add_icon_images(doc, icons)
            
            doc.add_paragraph('')
            print(f"   ✅ Found {len(icons)} icons: {', '.join(icons[:3])}{'...' if len(icons) > 3 else ''}")
            
        except Exception as e:
            doc.add_paragraph(f'❌ Error: {str(e)}')
            print(f"   ❌ Error: {e}")
    
    # Save document
    output_path = Path('outputs/comprehensive_test_results.docx')
    output_path.parent.mkdir(exist_ok=True)
    doc.save(output_path)
    
    print("\n" + "=" * 60)
    print(f"✅ Test complete! Results saved to: {output_path}")
    print(f"📄 Open the Word document to review all 30 tests with images")

if __name__ == '__main__':
    main()
